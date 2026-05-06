import csv
import hashlib
import hmac
import io
import multiprocessing
import os
import re
import tempfile
import uuid
import zipfile

import ocrmypdf
import pandas as pd
import razorpay
from docx import Document
from flask import Flask, after_this_request, jsonify, render_template_string, request, send_file
from openpyxl import load_workbook
from PIL import Image, ImageEnhance, ImageFilter
from pypdf import PdfReader, PdfWriter
from werkzeug.utils import secure_filename


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024

PDF_PRICE_PER_PAGE_PAISE = 20
IMAGE_PRICE_PAISE = 1000
PENDING_ORDERS = {}
BYPASS_ACCESS_KEY = os.getenv("OCR_BYPASS_KEY", "215836").strip()

PDF_SERVICES = {
    "ocr": "PDF OCR",
    "pdf-to-word": "PDF to Word",
    "pdf-splitter": "PDF Splitter",
    "pdf-delete-pages": "PDF Delete Pages",
    "pdf-summary": "PDF Summary",
}
IMAGE_SERVICES = {
    "blur-to-hd": "Blur to HD Image",
    "background-remover": "Background Remover",
}


PAGE = """
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>SCANLY</title>
  <style>
    :root {
      color-scheme: dark;
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      color: #f7fbff;
      background: #071018;
      --panel: rgba(255, 255, 255, 0.08);
      --line: rgba(255, 255, 255, 0.16);
      --cyan: #43d9ff;
      --lime: #b9f85d;
      --coral: #ff7b6e;
      --ink: #071018;
    }
    * { box-sizing: border-box; }
    html { scroll-behavior: smooth; }
    body {
      margin: 0;
      min-height: 100vh;
      background:
        radial-gradient(circle at 12% 12%, rgba(67, 217, 255, 0.2), transparent 24rem),
        radial-gradient(circle at 88% 18%, rgba(185, 248, 93, 0.13), transparent 22rem),
        linear-gradient(135deg, #071018 0%, #10212d 46%, #142117 100%);
      overflow-x: hidden;
    }
    body::before {
      content: "";
      position: fixed;
      inset: 0;
      pointer-events: none;
      background-image:
        linear-gradient(rgba(255,255,255,0.045) 1px, transparent 1px),
        linear-gradient(90deg, rgba(255,255,255,0.045) 1px, transparent 1px);
      background-size: 48px 48px;
      mask-image: linear-gradient(to bottom, black, transparent 85%);
    }
    main {
      width: min(1180px, calc(100% - 32px));
      margin: 0 auto;
      position: relative;
      z-index: 1;
    }
    .nav {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 18px;
      min-height: 76px;
    }
    .brand {
      display: inline-flex;
      color: #ffffff;
      text-decoration: none;
      font-size: 1.4rem;
      font-weight: 950;
      transform-style: preserve-3d;
      text-shadow: 0 1px 0 #7dd3fc, 0 10px 22px rgba(67, 217, 255, 0.22);
      transition: transform 220ms ease, text-shadow 220ms ease, color 220ms ease;
    }
    .brand:hover {
      color: var(--lime);
      transform: perspective(700px) rotateX(14deg) rotateY(-12deg) translateY(-2px);
      text-shadow: 0 1px 0 #ffffff, 0 12px 30px rgba(185, 248, 93, 0.4);
    }
    .nav-links {
      display: flex;
      align-items: center;
      gap: 8px;
      flex-wrap: wrap;
      justify-content: flex-end;
    }
    .nav-links a {
      min-height: 38px;
      display: inline-flex;
      align-items: center;
      padding: 8px 12px;
      color: #d8e6ef;
      text-decoration: none;
      border: 1px solid transparent;
      border-radius: 6px;
      font-weight: 750;
      font-size: 0.92rem;
    }
    .nav-links a:hover {
      border-color: var(--line);
      background: rgba(255,255,255,0.08);
      color: #ffffff;
    }
    .hero {
      min-height: calc(100vh - 76px);
      display: grid;
      grid-template-columns: minmax(0, 1.02fr) minmax(320px, 0.98fr);
      gap: 36px;
      align-items: center;
      padding: 16px 0 70px;
    }
    .hero-copy {
      display: grid;
      gap: 24px;
      max-width: 690px;
    }
    h1 {
      margin: 0;
      font-size: clamp(3.2rem, 8vw, 7.6rem);
      line-height: 0.9;
      letter-spacing: 0;
      color: #ffffff;
      text-shadow:
        0 2px 0 #69d9ff,
        0 8px 0 rgba(67, 217, 255, 0.12),
        0 28px 55px rgba(0,0,0,0.5);
    }
    .hero-copy p {
      margin: 0;
      max-width: 610px;
      color: #c8d8e2;
      font-size: 1.08rem;
      line-height: 1.58;
    }
    .price-row {
      display: flex;
      gap: 10px;
      flex-wrap: wrap;
    }
    .price-pill {
      display: inline-flex;
      align-items: center;
      min-height: 38px;
      padding: 8px 12px;
      border-radius: 6px;
      border: 1px solid rgba(255,255,255,0.16);
      background: rgba(255,255,255,0.08);
      color: #eef8ff;
      font-weight: 850;
    }
    .service-buttons {
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 10px;
    }
    .service-link {
      min-height: 50px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 12px 14px;
      border: 1px solid rgba(255,255,255,0.16);
      border-radius: 8px;
      background: rgba(255,255,255,0.08);
      color: #ffffff;
      text-decoration: none;
      font-weight: 850;
      box-shadow: inset 0 1px 0 rgba(255,255,255,0.12);
      transition: transform 180ms ease, background 180ms ease, border-color 180ms ease;
    }
    .service-link:hover {
      transform: translateY(-3px);
      background: rgba(67,217,255,0.16);
      border-color: rgba(67,217,255,0.55);
    }
    .service-link span:last-child {
      color: var(--lime);
      font-size: 1.1rem;
    }
    .scene {
      min-height: 560px;
      position: relative;
      perspective: 1000px;
    }
    .orbital {
      position: absolute;
      inset: 6% 0 0 0;
      transform-style: preserve-3d;
      animation: floatScene 8s ease-in-out infinite;
    }
    .slab {
      position: absolute;
      border: 1px solid rgba(255,255,255,0.18);
      border-radius: 8px;
      background: linear-gradient(145deg, rgba(255,255,255,0.19), rgba(255,255,255,0.045));
      box-shadow: 0 28px 70px rgba(0,0,0,0.38), inset 0 1px 0 rgba(255,255,255,0.16);
      backdrop-filter: blur(14px);
      transform-style: preserve-3d;
    }
    .slab.one {
      width: 72%;
      height: 66%;
      left: 12%;
      top: 10%;
      transform: rotateY(-24deg) rotateX(13deg) translateZ(60px);
    }
    .slab.two {
      width: 44%;
      height: 34%;
      right: 1%;
      top: 2%;
      transform: rotateY(-32deg) rotateX(16deg) translateZ(140px);
    }
    .slab.three {
      width: 42%;
      height: 30%;
      left: 0;
      bottom: 10%;
      transform: rotateY(-16deg) rotateX(19deg) translateZ(110px);
    }
    .scan-beam {
      position: absolute;
      left: 18%;
      top: 34%;
      width: 70%;
      height: 14px;
      border-radius: 999px;
      background: linear-gradient(90deg, transparent, var(--cyan), var(--lime), transparent);
      box-shadow: 0 0 28px rgba(67,217,255,0.75);
      transform: rotateY(-24deg) rotateX(13deg) translateZ(170px);
      animation: scan 2.6s ease-in-out infinite;
    }
    .cube {
      position: absolute;
      width: 70px;
      aspect-ratio: 1;
      border: 1px solid rgba(185,248,93,0.38);
      border-radius: 8px;
      background: rgba(185,248,93,0.1);
      transform-style: preserve-3d;
      animation: cubeSpin 8s linear infinite;
    }
    .cube.a { right: 18%; bottom: 18%; }
    .cube.b { left: 10%; top: 18%; width: 46px; animation-duration: 6s; }
    @keyframes floatScene {
      0%, 100% { transform: translateY(0) rotateZ(0deg); }
      50% { transform: translateY(-18px) rotateZ(1deg); }
    }
    @keyframes scan {
      0%, 100% { top: 25%; opacity: 0.55; }
      50% { top: 66%; opacity: 1; }
    }
    @keyframes cubeSpin {
      from { transform: rotateX(0deg) rotateY(0deg); }
      to { transform: rotateX(360deg) rotateY(360deg); }
    }
    .message {
      margin: 0 0 20px;
      padding: 14px 16px;
      border: 1px solid rgba(255,123,110,0.45);
      border-radius: 8px;
      color: #ffe6e2;
      background: rgba(255,123,110,0.12);
    }
    .tools {
      padding: 76px 0 82px;
    }
    .section-head {
      display: grid;
      gap: 10px;
      margin-bottom: 22px;
    }
    h2 {
      margin: 0;
      color: #ffffff;
      font-size: clamp(1.8rem, 4vw, 3rem);
      letter-spacing: 0;
    }
    .section-head p {
      margin: 0;
      color: #b9c7d2;
      line-height: 1.5;
    }
    .tool-grid {
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 16px;
    }
    .tool-card {
      display: grid;
      gap: 16px;
      min-height: 430px;
      align-content: start;
      padding: 22px;
      border: 1px solid var(--line);
      border-radius: 8px;
      background: rgba(255,255,255,0.075);
      box-shadow: 0 22px 60px rgba(0,0,0,0.24);
      backdrop-filter: blur(14px);
    }
    .tool-top {
      display: grid;
      gap: 8px;
    }
    .tag {
      width: fit-content;
      padding: 6px 9px;
      border-radius: 5px;
      background: rgba(185,248,93,0.16);
      color: var(--lime);
      font-size: 0.78rem;
      font-weight: 900;
      text-transform: uppercase;
    }
    h3 {
      margin: 0;
      color: #ffffff;
      font-size: 1.28rem;
      letter-spacing: 0;
    }
    .tool-top p, .details {
      margin: 0;
      color: #b9c7d2;
      line-height: 1.5;
    }
    form {
      display: grid;
      gap: 13px;
    }
    input[type="file"],
    input[type="password"],
    input[type="text"],
    select {
      width: 100%;
      min-height: 46px;
      padding: 11px 13px;
      border: 1px solid rgba(255,255,255,0.18);
      border-radius: 6px;
      background: rgba(7,16,24,0.65);
      color: #ffffff;
      font: inherit;
    }
    input::placeholder { color: #89a0ae; }
    input[type="file"]::file-selector-button {
      margin-right: 10px;
      border: 0;
      border-radius: 5px;
      background: #d9f99d;
      color: #10210d;
      padding: 8px 11px;
      font-weight: 900;
    }
    button {
      min-height: 46px;
      border: 0;
      border-radius: 6px;
      background: linear-gradient(135deg, var(--cyan), var(--lime));
      color: #081018;
      font-weight: 950;
      cursor: pointer;
      font: inherit;
      box-shadow: 0 14px 28px rgba(67,217,255,0.2);
      transition: transform 160ms ease, filter 160ms ease;
    }
    button:hover { transform: translateY(-2px); filter: brightness(1.03); }
    button:disabled {
      cursor: wait;
      background: #6b7c87;
      box-shadow: none;
      color: #e5eef4;
    }
    .billing-summary {
      display: none;
      gap: 6px;
      padding: 12px 14px;
      border: 1px solid rgba(67,217,255,0.34);
      border-radius: 6px;
      background: rgba(67,217,255,0.11);
      color: #effbff;
      font-weight: 850;
    }
    .billing-summary.ready { display: grid; }
    .billing-summary span:last-child {
      color: #a9bdc9;
      font-size: 0.9rem;
      font-weight: 650;
    }
    .progress-wrap {
      display: none;
      gap: 8px;
    }
    .progress-wrap.active { display: grid; }
    .progress-label {
      display: flex;
      justify-content: space-between;
      gap: 12px;
      color: #d6e5ee;
      font-size: 0.9rem;
      font-weight: 800;
    }
    .progress-track {
      height: 10px;
      overflow: hidden;
      border-radius: 999px;
      background: rgba(255,255,255,0.14);
    }
    .progress-bar {
      width: 0%;
      height: 100%;
      border-radius: inherit;
      background: linear-gradient(90deg, var(--cyan), var(--lime), var(--coral));
      transition: width 320ms ease;
    }
    .status {
      min-height: 24px;
      color: #bfefff;
      font-weight: 750;
    }
    .status.success { color: var(--lime); }
    .status.error { color: #ffb4ab; }
    .download-link {
      display: none;
      align-items: center;
      justify-content: center;
      min-height: 46px;
      border-radius: 6px;
      background: #ffffff;
      color: #071018;
      text-decoration: none;
      font-weight: 950;
    }
    .download-link.ready { display: flex; }
    .details {
      margin-top: auto;
      font-size: 0.92rem;
    }
    @media (max-width: 900px) {
      .hero {
        grid-template-columns: 1fr;
        padding-top: 28px;
      }
      .scene {
        min-height: 360px;
        order: -1;
      }
      .tool-grid {
        grid-template-columns: 1fr;
      }
    }
    @media (max-width: 620px) {
      main { width: min(100% - 22px, 1180px); }
      .nav { align-items: flex-start; flex-direction: column; padding-top: 14px; }
      .nav-links { justify-content: flex-start; }
      .service-buttons { grid-template-columns: 1fr; }
      .hero { min-height: auto; }
      .tool-card { min-height: auto; padding: 18px; }
    }
  </style>
</head>
<body>
  <main>
    <nav class="nav" aria-label="Main">
      <a class="brand" href="#">SCANLY</a>
      <div class="nav-links">
        <a href="#home">Home</a>
        <a href="#services">Services</a>
        <a href="#ocr">OCR</a>
        <a href="#images">Images</a>
      </div>
    </nav>

    <section class="hero" id="home">
      <div class="hero-copy">
        <h1>SCANLY</h1>
        <p>Convert, clean, split, summarize, and enhance files from one payment-ready workspace. Pick a service, upload the file, pay only after the page count is calculated, then download when processing is complete.</p>
        <div class="price-row">
          <span class="price-pill">PDF tools: INR 0.20 per page</span>
          <span class="price-pill">Image tools: INR 10.00 per image</span>
        </div>
        <div class="service-buttons" aria-label="SCANLY services">
          <a class="service-link" href="#ocr"><span>OCR</span><span>-></span></a>
          <a class="service-link" href="#excel"><span>Excel to CSV</span><span>-></span></a>
          <a class="service-link" href="#pdf-to-word"><span>PDF to Word</span><span>-></span></a>
          <a class="service-link" href="#pdf-splitter"><span>PDF Splitter</span><span>-></span></a>
          <a class="service-link" href="#pdf-delete-pages"><span>PDF Delete Page</span><span>-></span></a>
          <a class="service-link" href="#blur-to-hd"><span>Blur to HD Image</span><span>-></span></a>
          <a class="service-link" href="#background-remover"><span>Background Remover</span><span>-></span></a>
          <a class="service-link" href="#pdf-summary"><span>Summary for PDF</span><span>-></span></a>
        </div>
      </div>
      <div class="scene" aria-hidden="true">
        <div class="orbital">
          <div class="slab one"></div>
          <div class="slab two"></div>
          <div class="slab three"></div>
          <div class="scan-beam"></div>
          <div class="cube a"></div>
          <div class="cube b"></div>
        </div>
      </div>
    </section>

    {% if message %}
      <div class="message">{{ message }}</div>
    {% endif %}

    <section class="tools" id="services">
      <div class="section-head">
        <h2>Services</h2>
        <p>PDF services use Razorpay at INR 0.20 per page. Image services use INR 10.00 per image. Excel to CSV stays free.</p>
      </div>
      <div class="tool-grid">
        <article class="tool-card" id="ocr">
          <div class="tool-top">
            <span class="tag">PDF tool</span>
            <h3>OCR</h3>
            <p>Turn scanned PDFs into searchable documents. Fast mode skips pages that already contain selectable text.</p>
          </div>
          <form action="/ocr" method="post" enctype="multipart/form-data" data-download-form data-paid-form data-service="ocr" data-price-note="Rate: INR 0.20 per PDF page. A valid access key skips payment.">
            <input type="file" name="pdf_file" accept="application/pdf,.pdf" required>
            <select name="mode" aria-label="OCR mode">
              <option value="fast" selected>Fast OCR - skip pages that already have text</option>
              <option value="accurate">High accuracy - slower full OCR</option>
            </select>
            <input type="password" name="access_key" inputmode="numeric" autocomplete="off" placeholder="Access key for free processing">
            <button type="submit">Pay and Process</button>
            <div class="billing-summary" aria-live="polite"><span data-billing-total></span><span data-billing-note></span></div>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">Best for scanned reports, contracts, invoices, and office paperwork.</p>
        </article>

        <article class="tool-card" id="excel">
          <div class="tool-top">
            <span class="tag">Free tool</span>
            <h3>Excel to CSV</h3>
            <p>Convert the first worksheet into a clean UTF-8 CSV file.</p>
          </div>
          <form action="/excel" method="post" enctype="multipart/form-data" data-download-form>
            <input type="file" name="excel_file" accept=".xlsx,.xls" required>
            <button type="submit">Convert to CSV</button>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">Best for accounting sheets, exports, tables, and lead lists.</p>
        </article>

        <article class="tool-card" id="pdf-to-word">
          <div class="tool-top">
            <span class="tag">PDF tool</span>
            <h3>PDF to Word</h3>
            <p>Extract readable PDF text into a Word document.</p>
          </div>
          <form action="/pdf-to-word" method="post" enctype="multipart/form-data" data-download-form data-paid-form data-service="pdf-to-word" data-price-note="Rate: INR 0.20 per PDF page.">
            <input type="file" name="pdf_file" accept="application/pdf,.pdf" required>
            <input type="password" name="access_key" inputmode="numeric" autocomplete="off" placeholder="Access key for free processing">
            <button type="submit">Pay and Convert</button>
            <div class="billing-summary" aria-live="polite"><span data-billing-total></span><span data-billing-note></span></div>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">Works best on PDFs that already contain selectable text or have been OCR processed.</p>
        </article>

        <article class="tool-card" id="pdf-splitter">
          <div class="tool-top">
            <span class="tag">PDF tool</span>
            <h3>PDF Splitter</h3>
            <p>Split one PDF into separate single-page PDF files inside a ZIP.</p>
          </div>
          <form action="/pdf-splitter" method="post" enctype="multipart/form-data" data-download-form data-paid-form data-service="pdf-splitter" data-price-note="Rate: INR 0.20 per PDF page.">
            <input type="file" name="pdf_file" accept="application/pdf,.pdf" required>
            <input type="password" name="access_key" inputmode="numeric" autocomplete="off" placeholder="Access key for free processing">
            <button type="submit">Pay and Split</button>
            <div class="billing-summary" aria-live="polite"><span data-billing-total></span><span data-billing-note></span></div>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">Useful for separating forms, annexures, scanned packs, and page-wise records.</p>
        </article>

        <article class="tool-card" id="pdf-delete-pages">
          <div class="tool-top">
            <span class="tag">PDF tool</span>
            <h3>PDF Delete Page</h3>
            <p>Remove selected pages and download a cleaned PDF.</p>
          </div>
          <form action="/pdf-delete-pages" method="post" enctype="multipart/form-data" data-download-form data-paid-form data-service="pdf-delete-pages" data-price-note="Rate: INR 0.20 per PDF page.">
            <input type="file" name="pdf_file" accept="application/pdf,.pdf" required>
            <input type="text" name="pages_to_delete" placeholder="Pages to delete, example: 1,3-5" required>
            <input type="password" name="access_key" inputmode="numeric" autocomplete="off" placeholder="Access key for free processing">
            <button type="submit">Pay and Delete Pages</button>
            <div class="billing-summary" aria-live="polite"><span data-billing-total></span><span data-billing-note></span></div>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">Page numbers are 1-based. Ranges are supported.</p>
        </article>

        <article class="tool-card" id="pdf-summary">
          <div class="tool-top">
            <span class="tag">PDF tool</span>
            <h3>Summary for PDF</h3>
            <p>Create a short text summary from extracted PDF text.</p>
          </div>
          <form action="/pdf-summary" method="post" enctype="multipart/form-data" data-download-form data-paid-form data-service="pdf-summary" data-price-note="Rate: INR 0.20 per PDF page.">
            <input type="file" name="pdf_file" accept="application/pdf,.pdf" required>
            <input type="password" name="access_key" inputmode="numeric" autocomplete="off" placeholder="Access key for free processing">
            <button type="submit">Pay and Summarize</button>
            <div class="billing-summary" aria-live="polite"><span data-billing-total></span><span data-billing-note></span></div>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">For scanned PDFs, run OCR first, then summarize the searchable PDF.</p>
        </article>

        <article class="tool-card" id="images">
          <div class="tool-top">
            <span class="tag">Image tools</span>
            <h3>Image Processing</h3>
            <p>Use local image filters for sharper output or a simple background cutout.</p>
          </div>
          <p class="details">Each image tool costs INR 10.00 per image.</p>
        </article>

        <article class="tool-card" id="blur-to-hd">
          <div class="tool-top">
            <span class="tag">Image tool</span>
            <h3>Blur to HD Image</h3>
            <p>Upscale, sharpen, and improve contrast for soft images.</p>
          </div>
          <form action="/blur-to-hd" method="post" enctype="multipart/form-data" data-download-form data-paid-form data-service="blur-to-hd" data-price-note="Rate: INR 10.00 per image.">
            <input type="file" name="image_file" accept="image/png,image/jpeg,image/webp,.png,.jpg,.jpeg,.webp" required>
            <input type="password" name="access_key" inputmode="numeric" autocomplete="off" placeholder="Access key for free processing">
            <button type="submit">Pay and Enhance</button>
            <div class="billing-summary" aria-live="polite"><span data-billing-total></span><span data-billing-note></span></div>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">Best for mildly blurred screenshots, IDs, and document photos.</p>
        </article>

        <article class="tool-card" id="background-remover">
          <div class="tool-top">
            <span class="tag">Image tool</span>
            <h3>Background Remover</h3>
            <p>Remove a plain background by detecting the image corner color.</p>
          </div>
          <form action="/background-remover" method="post" enctype="multipart/form-data" data-download-form data-paid-form data-service="background-remover" data-price-note="Rate: INR 10.00 per image.">
            <input type="file" name="image_file" accept="image/png,image/jpeg,image/webp,.png,.jpg,.jpeg,.webp" required>
            <input type="password" name="access_key" inputmode="numeric" autocomplete="off" placeholder="Access key for free processing">
            <button type="submit">Pay and Remove Background</button>
            <div class="billing-summary" aria-live="polite"><span data-billing-total></span><span data-billing-note></span></div>
            <div class="progress-wrap" aria-hidden="true"><div class="progress-label"><span>Processing</span><span data-progress-value>0%</span></div><div class="progress-track"><div class="progress-bar"></div></div></div>
            <div class="status" aria-live="polite"></div>
            <a class="download-link" href="#" download>Download result</a>
          </form>
          <p class="details">Works best with white, light, or single-color backgrounds.</p>
        </article>
      </div>
    </section>
  </main>
  <script src="https://checkout.razorpay.com/v1/checkout.js"></script>
  <script>
    function filenameFromDisposition(disposition, fallback) {
      if (!disposition) return fallback;
      const match = disposition.match(/filename="?([^"]+)"?/i);
      return match ? match[1] : fallback;
    }

    async function errorFromResponse(response, fallback) {
      const contentType = response.headers.get("content-type") || "";
      if (contentType.includes("application/json")) {
        const payload = await response.json();
        return payload.error || fallback;
      }
      const message = await response.text();
      return message.replace(/<[^>]*>/g, " ").replace(/\\s+/g, " ").trim() || fallback;
    }

    async function createPaymentOrder(form) {
      const data = new FormData(form);
      data.append("service", form.dataset.service);
      const response = await fetch("/payment/order", {
        method: "POST",
        body: data,
      });
      if (!response.ok) {
        throw new Error(await errorFromResponse(response, "Could not create payment order."));
      }
      return response.json();
    }

    function openRazorpayCheckout(order) {
      return new Promise((resolve, reject) => {
        if (!window.Razorpay) {
          reject(new Error("Razorpay checkout could not load. Please refresh and try again."));
          return;
        }
        const checkout = new Razorpay({
          key: order.key,
          amount: order.amount,
          currency: order.currency,
          name: "SCANLY",
          description: order.description,
          order_id: order.order_id,
          handler: resolve,
          modal: {
            ondismiss: () => reject(new Error("Payment was not completed.")),
          },
          theme: {
            color: "#43d9ff",
          },
        });
        checkout.on("payment.failed", (response) => {
          reject(new Error(response.error && response.error.description ? response.error.description : "Payment failed."));
        });
        checkout.open();
      });
    }

    document.querySelectorAll("[data-download-form]").forEach((form) => {
      let activeDownloadUrl = null;

      form.addEventListener("submit", async (event) => {
        event.preventDefault();

        const button = form.querySelector("button");
        const status = form.querySelector(".status");
        const progressWrap = form.querySelector(".progress-wrap");
        const progressBar = form.querySelector(".progress-bar");
        const progressValue = form.querySelector("[data-progress-value]");
        const downloadLink = form.querySelector(".download-link");
        const billingSummary = form.querySelector(".billing-summary");
        const billingTotal = form.querySelector("[data-billing-total]");
        const billingNote = form.querySelector("[data-billing-note]");
        const originalText = button.textContent;
        const requiresPayment = form.hasAttribute("data-paid-form");
        const fallbackName = form.action.endsWith("/excel") ? "converted.csv" : "scanly-output";
        let progress = 0;
        let progressTimer = null;

        if (activeDownloadUrl) {
          URL.revokeObjectURL(activeDownloadUrl);
          activeDownloadUrl = null;
        }
        downloadLink.classList.remove("ready");
        downloadLink.removeAttribute("href");
        if (billingSummary) billingSummary.classList.remove("ready");

        const setProgress = (value) => {
          progress = Math.max(progress, Math.min(value, 100));
          progressBar.style.width = `${progress}%`;
          progressValue.textContent = `${Math.round(progress)}%`;
        };

        button.disabled = true;
        button.textContent = requiresPayment ? "Calculating..." : "Processing...";
        progressWrap.classList.add("active");
        progressWrap.setAttribute("aria-hidden", "false");
        setProgress(8);
        status.classList.remove("error", "success");
        status.textContent = requiresPayment ? "Uploading file to calculate payment." : "Uploading and converting.";

        try {
          const formData = new FormData(form);
          if (requiresPayment) {
            const accessKey = formData.get("access_key");
            if (accessKey && accessKey.trim()) {
              if (billingSummary && billingTotal && billingNote) {
                billingTotal.textContent = "Access key entered. Payment will be skipped after server verification.";
                billingNote.textContent = form.dataset.priceNote || "";
                billingSummary.classList.add("ready");
              }
              setProgress(20);
              button.textContent = "Processing...";
              status.textContent = "Checking access key and starting service.";
            } else {
              const order = await createPaymentOrder(form);
              if (billingSummary && billingTotal && billingNote) {
                billingTotal.textContent = order.summary;
                billingNote.textContent = form.dataset.priceNote || "";
                billingSummary.classList.add("ready");
              }
              setProgress(18);
              button.textContent = "Waiting for payment...";
              status.textContent = `Complete INR ${order.display_amount} payment to start.`;
              const payment = await openRazorpayCheckout(order);
              formData.append("razorpay_payment_id", payment.razorpay_payment_id);
              formData.append("razorpay_order_id", payment.razorpay_order_id);
              formData.append("razorpay_signature", payment.razorpay_signature);
              setProgress(26);
              button.textContent = "Processing...";
              status.textContent = "Payment verified. Processing now.";
            }
          }

          progressTimer = window.setInterval(() => {
            const next = progress < 55 ? progress + 7 : progress < 86 ? progress + 3 : progress + 0.6;
            setProgress(Math.min(next, 94));
          }, 650);

          const response = await fetch(form.action, {
            method: "POST",
            body: formData,
          });
          if (!response.ok) {
            throw new Error(await errorFromResponse(response, "Processing failed."));
          }
          const blob = await response.blob();
          activeDownloadUrl = URL.createObjectURL(blob);
          downloadLink.href = activeDownloadUrl;
          downloadLink.download = filenameFromDisposition(response.headers.get("content-disposition"), fallbackName);
          downloadLink.classList.add("ready");
          setProgress(100);
          status.classList.add("success");
          status.textContent = "Processed. Your file is ready to download.";
        } catch (error) {
          setProgress(100);
          status.classList.add("error");
          status.textContent = error.message || "Processing failed.";
        } finally {
          if (progressTimer) window.clearInterval(progressTimer);
          button.disabled = false;
          button.textContent = originalText;
        }
      });
    });
  </script>
</body>
</html>
"""


def render_page(message=None):
    return render_template_string(PAGE, message=message)


@app.get("/")
def index():
    return render_page()


@app.get("/health")
def health():
    return {"status": "ok"}


def remove_later(path):
    def cleanup():
        try:
            os.remove(path)
        except OSError:
            pass

    @after_this_request
    def schedule_cleanup(response):
        response.call_on_close(cleanup)
        return response


def razorpay_keys():
    key_id = os.getenv("RAZORPAY_KEY_ID", "").strip()
    key_secret = os.getenv("RAZORPAY_KEY_SECRET", "").strip()
    if not key_id or not key_secret:
        raise RuntimeError("Razorpay keys are not configured. Set RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET.")
    return key_id, key_secret


def verify_razorpay_signature(order_id, payment_id, signature):
    _, key_secret = razorpay_keys()
    message = f"{order_id}|{payment_id}".encode("utf-8")
    expected = hmac.new(key_secret.encode("utf-8"), message, hashlib.sha256).hexdigest()
    return hmac.compare_digest(expected, signature or "")


def has_valid_bypass_key():
    access_key = request.form.get("access_key", "").strip()
    return bool(BYPASS_ACCESS_KEY and hmac.compare_digest(access_key, BYPASS_ACCESS_KEY))


def count_pdf_pages(path):
    reader = PdfReader(path)
    return len(reader.pages)


def save_pdf_upload(uploaded):
    if not uploaded or uploaded.filename == "":
        raise ValueError("Please upload a PDF file.")
    filename = secure_filename(uploaded.filename)
    if not filename.lower().endswith(".pdf"):
        raise ValueError("Please upload a valid PDF file.")
    input_file = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    input_file.close()
    uploaded.save(input_file.name)
    return filename, input_file.name


def save_image_upload(uploaded):
    if not uploaded or uploaded.filename == "":
        raise ValueError("Please upload an image file.")
    filename = secure_filename(uploaded.filename)
    extension = os.path.splitext(filename)[1].lower()
    if extension not in {".png", ".jpg", ".jpeg", ".webp"}:
        raise ValueError("Please upload a PNG, JPG, JPEG, or WEBP image.")
    input_file = tempfile.NamedTemporaryFile(suffix=extension, delete=False)
    input_file.close()
    uploaded.save(input_file.name)
    return filename, input_file.name


def order_amount_for(service, path):
    if service in PDF_SERVICES:
        pages = count_pdf_pages(path)
        if pages < 1:
            raise ValueError("This PDF does not contain any pages.")
        amount = pages * PDF_PRICE_PER_PAGE_PAISE
        return amount, pages, f"{pages} page{'s' if pages != 1 else ''} x INR 0.20 = INR {amount / 100:.2f}"
    if service in IMAGE_SERVICES:
        return IMAGE_PRICE_PAISE, 1, "1 image x INR 10.00 = INR 10.00"
    raise ValueError("Unknown paid service.")


def create_order(service, amount, pages):
    key_id, key_secret = razorpay_keys()
    client = razorpay.Client(auth=(key_id, key_secret))
    order = client.order.create(
        {
            "amount": amount,
            "currency": "INR",
            "receipt": f"scanly_{uuid.uuid4().hex[:25]}",
            "notes": {
                "service": service,
                "pages": str(pages),
                "amount": str(amount),
            },
        }
    )
    PENDING_ORDERS[order["id"]] = {
        "amount": amount,
        "pages": pages,
        "service": service,
        "currency": "INR",
    }
    return key_id, order


@app.post("/payment/order")
def create_payment_order():
    service = request.form.get("service", "").strip()
    input_path = None
    try:
        if service in PDF_SERVICES:
            _, input_path = save_pdf_upload(request.files.get("pdf_file"))
        elif service in IMAGE_SERVICES:
            _, input_path = save_image_upload(request.files.get("image_file"))
            Image.open(input_path).verify()
        else:
            return jsonify({"error": "Unknown service."}), 400

        amount, pages, summary = order_amount_for(service, input_path)
        key_id, order = create_order(service, amount, pages)
        description = PDF_SERVICES.get(service) or IMAGE_SERVICES.get(service) or "SCANLY service"
        return jsonify(
            {
                "key": key_id,
                "order_id": order["id"],
                "amount": amount,
                "currency": "INR",
                "pages": pages,
                "display_amount": f"{amount / 100:.2f}",
                "summary": summary,
                "description": description,
            }
        )
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"Could not create payment order: {exc}"}), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass


@app.post("/ocr/order")
def create_ocr_order():
    input_path = None
    try:
        _, input_path = save_pdf_upload(request.files.get("pdf_file"))
        amount, pages, summary = order_amount_for("ocr", input_path)
        key_id, order = create_order("ocr", amount, pages)
        return jsonify(
            {
                "key": key_id,
                "order_id": order["id"],
                "amount": amount,
                "currency": "INR",
                "pages": pages,
                "display_amount": f"{amount / 100:.2f}",
                "summary": summary,
                "description": "PDF OCR",
            }
        )
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"Could not create payment order: {exc}"}), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass


def payment_error_for(service, amount, pages):
    if has_valid_bypass_key():
        return None
    order_id = request.form.get("razorpay_order_id", "")
    payment_id = request.form.get("razorpay_payment_id", "")
    signature = request.form.get("razorpay_signature", "")
    order = PENDING_ORDERS.get(order_id)
    if not order_id or not payment_id or not signature:
        return render_page("Complete the Razorpay payment before processing."), 402
    if not order:
        return render_page("Payment order expired. Please calculate payment and try again."), 402
    if order["service"] != service or order["amount"] != amount or order["pages"] != pages:
        return render_page("Uploaded file changed after payment. Please pay again for this file."), 400
    if not verify_razorpay_signature(order_id, payment_id, signature):
        return render_page("Payment verification failed. Please try again."), 402
    return None


def finish_paid_order():
    order_id = request.form.get("razorpay_order_id", "")
    if order_id:
        PENDING_ORDERS.pop(order_id, None)


def ocr_options(mode):
    workers = max(1, min(4, multiprocessing.cpu_count()))
    if mode == "accurate":
        return {
            "force_ocr": True,
            "deskew": True,
            "rotate_pages": True,
            "oversample": 300,
            "optimize": 1,
            "jobs": workers,
            "progress_bar": False,
        }
    return {
        "skip_text": True,
        "deskew": False,
        "rotate_pages": False,
        "oversample": 150,
        "optimize": 1,
        "jobs": workers,
        "progress_bar": False,
    }


@app.post("/ocr")
def ocr_pdf():
    output_file = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    output_file.close()
    output_path = output_file.name
    output_sent = False
    input_path = None
    try:
        filename, input_path = save_pdf_upload(request.files.get("pdf_file"))
        amount, pages, _ = order_amount_for("ocr", input_path)
        payment_error = payment_error_for("ocr", amount, pages)
        if payment_error:
            return payment_error
        os.remove(output_path)
        ocrmypdf.ocr(input_path, output_path, **ocr_options(request.form.get("mode", "fast")))
        remove_later(output_path)
        finish_paid_order()
        output_sent = True
        return send_file(output_path, as_attachment=True, download_name=f"scanly_ocr_{filename}")
    except ValueError as exc:
        return render_page(str(exc)), 400
    except Exception as exc:
        return render_page(f"OCR failed: {exc}"), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass
        if not output_sent:
            try:
                os.remove(output_path)
            except OSError:
                pass


def extract_pdf_text(reader):
    chunks = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(f"Page {index}\n{text.strip()}")
    return "\n\n".join(chunks).strip()


@app.post("/pdf-to-word")
def pdf_to_word():
    input_path = None
    output_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    output_file.close()
    output_sent = False
    try:
        filename, input_path = save_pdf_upload(request.files.get("pdf_file"))
        amount, pages, _ = order_amount_for("pdf-to-word", input_path)
        payment_error = payment_error_for("pdf-to-word", amount, pages)
        if payment_error:
            return payment_error
        reader = PdfReader(input_path)
        doc = Document()
        doc.add_heading("SCANLY PDF to Word", 0)
        for index, page in enumerate(reader.pages, start=1):
            doc.add_heading(f"Page {index}", level=1)
            text = page.extract_text() or ""
            doc.add_paragraph(text.strip() or "[No selectable text found on this page.]")
        doc.save(output_file.name)
        remove_later(output_file.name)
        finish_paid_order()
        output_sent = True
        download_name = f"{os.path.splitext(filename)[0]}.docx"
        return send_file(output_file.name, as_attachment=True, download_name=download_name)
    except ValueError as exc:
        return render_page(str(exc)), 400
    except Exception as exc:
        return render_page(f"PDF to Word failed: {exc}"), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass
        if not output_sent:
            try:
                os.remove(output_file.name)
            except OSError:
                pass


@app.post("/pdf-splitter")
def pdf_splitter():
    input_path = None
    output_file = tempfile.NamedTemporaryFile(suffix=".zip", delete=False)
    output_file.close()
    output_sent = False
    try:
        filename, input_path = save_pdf_upload(request.files.get("pdf_file"))
        amount, pages, _ = order_amount_for("pdf-splitter", input_path)
        payment_error = payment_error_for("pdf-splitter", amount, pages)
        if payment_error:
            return payment_error
        reader = PdfReader(input_path)
        stem = os.path.splitext(filename)[0]
        with zipfile.ZipFile(output_file.name, "w", zipfile.ZIP_DEFLATED) as archive:
            for index, page in enumerate(reader.pages, start=1):
                writer = PdfWriter()
                writer.add_page(page)
                buffer = io.BytesIO()
                writer.write(buffer)
                archive.writestr(f"{stem}_page_{index}.pdf", buffer.getvalue())
        remove_later(output_file.name)
        finish_paid_order()
        output_sent = True
        return send_file(output_file.name, as_attachment=True, download_name=f"{stem}_split_pages.zip")
    except ValueError as exc:
        return render_page(str(exc)), 400
    except Exception as exc:
        return render_page(f"PDF split failed: {exc}"), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass
        if not output_sent:
            try:
                os.remove(output_file.name)
            except OSError:
                pass


def parse_pages_to_delete(value, max_pages):
    pages = set()
    for part in value.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start_text, end_text = part.split("-", 1)
            start = int(start_text.strip())
            end = int(end_text.strip())
            if start > end:
                start, end = end, start
            pages.update(range(start, end + 1))
        else:
            pages.add(int(part))
    if not pages:
        raise ValueError("Enter at least one page to delete.")
    invalid = [page for page in pages if page < 1 or page > max_pages]
    if invalid:
        raise ValueError(f"Page number out of range: {invalid[0]}.")
    if len(pages) >= max_pages:
        raise ValueError("You cannot delete every page from a PDF.")
    return pages


@app.post("/pdf-delete-pages")
def pdf_delete_pages():
    input_path = None
    output_file = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    output_file.close()
    output_sent = False
    try:
        filename, input_path = save_pdf_upload(request.files.get("pdf_file"))
        amount, pages, _ = order_amount_for("pdf-delete-pages", input_path)
        payment_error = payment_error_for("pdf-delete-pages", amount, pages)
        if payment_error:
            return payment_error
        reader = PdfReader(input_path)
        delete_pages = parse_pages_to_delete(request.form.get("pages_to_delete", ""), len(reader.pages))
        writer = PdfWriter()
        for index, page in enumerate(reader.pages, start=1):
            if index not in delete_pages:
                writer.add_page(page)
        with open(output_file.name, "wb") as output:
            writer.write(output)
        remove_later(output_file.name)
        finish_paid_order()
        output_sent = True
        return send_file(output_file.name, as_attachment=True, download_name=f"cleaned_{filename}")
    except ValueError as exc:
        return render_page(str(exc)), 400
    except Exception as exc:
        return render_page(f"PDF page deletion failed: {exc}"), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass
        if not output_sent:
            try:
                os.remove(output_file.name)
            except OSError:
                pass


def summarize_text(text, limit=6):
    clean = re.sub(r"\s+", " ", text).strip()
    if not clean:
        return "No selectable text was found. Run OCR first, then summarize the searchable PDF."
    sentences = re.split(r"(?<=[.!?])\s+", clean)
    words = re.findall(r"[A-Za-z]{4,}", clean.lower())
    stop = {"that", "this", "with", "from", "have", "were", "their", "there", "which", "would", "could", "should", "about", "after", "before", "into", "than"}
    freq = {}
    for word in words:
        if word not in stop:
            freq[word] = freq.get(word, 0) + 1
    ranked = []
    for index, sentence in enumerate(sentences):
        score = sum(freq.get(word, 0) for word in re.findall(r"[A-Za-z]{4,}", sentence.lower()))
        if len(sentence) > 280:
            score *= 0.65
        ranked.append((score, index, sentence))
    chosen = sorted(sorted(ranked, reverse=True)[:limit], key=lambda item: item[1])
    return "\n".join(f"- {sentence.strip()}" for _, _, sentence in chosen if sentence.strip())


@app.post("/pdf-summary")
def pdf_summary():
    input_path = None
    output_file = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8")
    output_path = output_file.name
    output_file.close()
    output_sent = False
    try:
        filename, input_path = save_pdf_upload(request.files.get("pdf_file"))
        amount, pages, _ = order_amount_for("pdf-summary", input_path)
        payment_error = payment_error_for("pdf-summary", amount, pages)
        if payment_error:
            return payment_error
        reader = PdfReader(input_path)
        summary = summarize_text(extract_pdf_text(reader))
        with open(output_path, "w", encoding="utf-8") as output:
            output.write("SCANLY PDF Summary\n\n")
            output.write(summary)
            output.write("\n")
        remove_later(output_path)
        finish_paid_order()
        output_sent = True
        return send_file(output_path, as_attachment=True, download_name=f"{os.path.splitext(filename)[0]}_summary.txt")
    except ValueError as exc:
        return render_page(str(exc)), 400
    except Exception as exc:
        return render_page(f"PDF summary failed: {exc}"), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass
        if not output_sent:
            try:
                os.remove(output_path)
            except OSError:
                pass


@app.post("/blur-to-hd")
def blur_to_hd():
    input_path = None
    output_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    output_file.close()
    output_sent = False
    try:
        filename, input_path = save_image_upload(request.files.get("image_file"))
        amount, pages, _ = order_amount_for("blur-to-hd", input_path)
        payment_error = payment_error_for("blur-to-hd", amount, pages)
        if payment_error:
            return payment_error
        image = Image.open(input_path).convert("RGB")
        width, height = image.size
        image = image.resize((width * 2, height * 2), Image.Resampling.LANCZOS)
        image = ImageEnhance.Sharpness(image).enhance(2.4)
        image = ImageEnhance.Contrast(image).enhance(1.12)
        image = image.filter(ImageFilter.UnsharpMask(radius=2, percent=160, threshold=3))
        image.save(output_file.name, "PNG", optimize=True)
        remove_later(output_file.name)
        finish_paid_order()
        output_sent = True
        return send_file(output_file.name, as_attachment=True, download_name=f"hd_{os.path.splitext(filename)[0]}.png")
    except ValueError as exc:
        return render_page(str(exc)), 400
    except Exception as exc:
        return render_page(f"Image enhancement failed: {exc}"), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass
        if not output_sent:
            try:
                os.remove(output_file.name)
            except OSError:
                pass


def remove_plain_background(image):
    image = image.convert("RGBA")
    width, height = image.size
    pixels = image.load()
    sample_points = [(0, 0), (width - 1, 0), (0, height - 1), (width - 1, height - 1)]
    samples = [pixels[x, y][:3] for x, y in sample_points]
    background = tuple(sum(color[channel] for color in samples) // len(samples) for channel in range(3))
    threshold = 46
    for y in range(height):
        for x in range(width):
            red, green, blue, alpha = pixels[x, y]
            distance = abs(red - background[0]) + abs(green - background[1]) + abs(blue - background[2])
            if distance < threshold:
                pixels[x, y] = (red, green, blue, 0)
            elif distance < threshold * 2:
                pixels[x, y] = (red, green, blue, max(80, alpha - 80))
    return image


@app.post("/background-remover")
def background_remover():
    input_path = None
    output_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    output_file.close()
    output_sent = False
    try:
        filename, input_path = save_image_upload(request.files.get("image_file"))
        amount, pages, _ = order_amount_for("background-remover", input_path)
        payment_error = payment_error_for("background-remover", amount, pages)
        if payment_error:
            return payment_error
        image = Image.open(input_path)
        output = remove_plain_background(image)
        output.save(output_file.name, "PNG", optimize=True)
        remove_later(output_file.name)
        finish_paid_order()
        output_sent = True
        return send_file(output_file.name, as_attachment=True, download_name=f"transparent_{os.path.splitext(filename)[0]}.png")
    except ValueError as exc:
        return render_page(str(exc)), 400
    except Exception as exc:
        return render_page(f"Background removal failed: {exc}"), 500
    finally:
        if input_path:
            try:
                os.remove(input_path)
            except OSError:
                pass
        if not output_sent:
            try:
                os.remove(output_file.name)
            except OSError:
                pass


def convert_xlsx_to_csv(input_path, output_path):
    workbook = load_workbook(input_path, read_only=True, data_only=True)
    try:
        sheet = workbook.active
        with open(output_path, "w", newline="", encoding="utf-8") as csv_file:
            writer = csv.writer(csv_file)
            for row in sheet.iter_rows(values_only=True):
                writer.writerow(["" if value is None else value for value in row])
    finally:
        workbook.close()


def convert_excel_file(input_path, output_path, extension):
    if extension == ".xlsx":
        convert_xlsx_to_csv(input_path, output_path)
        return
    df = pd.read_excel(input_path, dtype=str, parse_dates=False, keep_default_na=False)
    df.to_csv(output_path, index=False, encoding="utf-8")


@app.post("/excel")
def excel_to_csv():
    uploaded = request.files.get("excel_file")
    if not uploaded or uploaded.filename == "":
        return render_page("Please upload an Excel file."), 400
    filename = secure_filename(uploaded.filename)
    if not filename.lower().endswith((".xlsx", ".xls")):
        return render_page("Please upload a valid Excel file."), 400
    extension = os.path.splitext(filename)[1].lower()
    input_file = tempfile.NamedTemporaryFile(suffix=extension, delete=False)
    output_file = tempfile.NamedTemporaryFile(suffix=".csv", delete=False)
    input_file.close()
    output_file.close()
    try:
        uploaded.save(input_file.name)
        convert_excel_file(input_file.name, output_file.name, extension)
        download_name = f"{os.path.splitext(filename)[0]}.csv"
        remove_later(output_file.name)
        return send_file(output_file.name, as_attachment=True, download_name=download_name)
    except Exception as exc:
        return render_page(f"Conversion failed: {exc}"), 500
    finally:
        try:
            os.remove(input_file.name)
        except OSError:
            pass


if __name__ == "__main__":
    app.run(host=os.getenv("HOST", "0.0.0.0"), port=int(os.getenv("PORT", "7860")))
